import streamlit as st
import os
import pandas as pd
import requests
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re
from streamlit_lottie import st_lottie

# Function to sanitize the property description for use in file names
def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|,]', '', filename)

# Function to create the prompt based on the Maharashtra-specific format
def create_prompt_maharashtra(date, ownership_history, missing_links):
    ownership_str = ""
    for i, owner in enumerate(ownership_history):
        if i == 0:
            ownership_str += f"{owner['owner_name']} (Acquired on: {owner['acquisition_date']}), Survey No.: {owner['survey_number']}, Area: {owner['area']}, Location: {owner['property_location']}"
        else:
            ownership_str += f"{owner['owner_name']} (Acquired on: {owner['acquisition_date']}, Sale Price: {owner['sale_price']})"
        if i < len(ownership_history) - 1:
            ownership_str += " → "

    current_owner = ownership_history[-1]['owner_name']

    return f"""
    Generate a Search Title Report (STR) based on the following details:

    Date: {date}

    1. Introduction:

    This Search Title Report has been prepared to provide a detailed examination of the title of the property, under the jurisdiction of the relevant local authority. The report covers the verification of ownership, encumbrances, and compliance with local regulations as of the date mentioned above.

    2. Legal Description of the Property:

    The property in question is identified by the history provided below. The legal description includes all boundaries and measurements as recorded in the land records.

    3. Chain of Title and Ownership History:

    The property has been traced back over the past 30 years, with the following owners listed in sequence: {ownership_str}. The current owner, {current_owner}, holds the title as per the latest transaction.

    4. Missing Links in Ownership Chain:

    {', '.join(missing_links) if missing_links else 'No missing links detected in the ownership chain.'}

    5. Compliance with Local Regulations:

    The property complies with the zoning regulations as specified by the relevant local authority. The necessary No Objection Certificates (NOCs) have been obtained from the relevant authorities. No government notifications or orders affecting the property were found during the search.

    6. Observations and Recommendations:

    During the search, the following observations were made: [Insert Any Irregularities, Concerns, or Potential Issues]. It is recommended that the purchaser conduct further due diligence, particularly in areas where discrepancies were noted. The report highlights the need for a professional title examination to confirm the findings before proceeding with any real estate transactions.

    7. Conclusion:

    This report is issued based on the available records as of the date mentioned above. While all efforts have been made to ensure accuracy, this report does not constitute a guarantee of title. The purchaser is advised to seek independent legal advice and to consider obtaining title insurance as an additional safeguard.

    Disclaimer:
    This Search Title Report is intended for informational purposes only and should not be construed as legal advice. It reflects the state of the title as per the records examined on {date}. The author of this report shall not be liable for any actions taken based on this information without further verification.

     Adv.Pravin D Sangvikar
    """

# Function to generate the report using the Google API
def generate_report(prompt, api_key):
    endpoint = 'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent'
    headers = {'Content-Type': 'application/json'}
    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    url = f"{endpoint}?key={api_key}"
    
    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()  # Raise an exception for HTTP errors
        response_data = response.json()
        content = response_data.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
        return content
    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error occurred: {http_err}")
    except Exception as err:
        st.error(f"Other error occurred: {err}")
    return None

# Function to convert the document into bytes for download
def convert_to_bytes(doc):
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

# Function to save the report content to a Word document and allow downloading
def save_to_word(report_content, file_name, advocate_name, advocate_address):
    document = Document()

    section = document.sections[0]
    header = section.header

    header_paragraph = header.paragraphs[0]
    header_paragraph.text = advocate_name
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header_paragraph2 = header.add_paragraph(advocate_address)
    header_paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('Search Title Report', level=1)
    document.add_paragraph(report_content)

    # Save the document to a BytesIO object
    file_bytes = convert_to_bytes(document)

    st.success(f"Search Title Report generated.")

    # Provide download link for the Word document
    st.download_button(
        label="Download Search Title Report",
        data=file_bytes,
        file_name=f"{file_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Function to load Lottie animation from URL
def load_lottie_url(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# Load the Lottie animation
lottie_animation_url = "https://lottie.host/c1b82357-ac83-4d42-95c7-931d7c7c8584/GRhvzcNSys.json"
lottie_animation = load_lottie_url(lottie_animation_url)

# Streamlit UI
st.title("Smart Search Title Report Generator")

# Display the Lottie animation
st_lottie(
    lottie_animation,
    speed=1,
    reverse=False,
    loop=True,
    quality="high",
    height=300,
    width=300,
    key="animation",
)

# Sidebar customization
st.sidebar.markdown(
    """
    <style>
    .sidebar .sidebar-content {
        background-color: #1E3A8A;  /* Ivory background */
    }
    </style>
    """,
    unsafe_allow_html=True
)
st.sidebar.title("Smart Search Title Report Generator")

# Uploading Excel files
uploaded_files = st.file_uploader("Upload Excel Files", type=["xlsx"], accept_multiple_files=True)

# User input
property_description = st.text_input("Property Description")
str_date = st.text_input("STR Generation Date (e.g., 01/01/2024)")

if st.button("Analyze"):
    if not uploaded_files:
        st.error("Please upload at least one Excel file.")
    elif not property_description or not str_date:
        st.error("Please provide both property description and generation date.")
    else:
        dataframes = []
        for uploaded_file in uploaded_files:
            try:
                df = pd.read_excel(uploaded_file, engine='openpyxl')
                expected_columns = [
                    'srocode', 'internaldocumentnumber', 'docno', 'docname', 'registrationdate',
                    'sroname', 'sellerparty', 'purchaserparty', 'propertydescription', 'areaname',
                    'consideration_amt', 'MarketValue', 'DateOfExecution', 'StampDutyPaid',
                    'RegistrationFees', 'status', 'micrno', 'party_code', 'bank_type'
                ]
                if all(col in df.columns for col in expected_columns):
                    dataframes.append(df[expected_columns])
                else:
                    st.warning(f"File {uploaded_file.name} does not have the required columns and will be skipped.")
            except Exception as e:
                st.error(f"Error reading {uploaded_file.name}: {str(e)}")

        results = []
        for df in dataframes:
            result = df[df['propertydescription'].str.contains(property_description, case=False, na=False, regex=False)]
            if not result.empty:
                results.append(result)

        if results:
            final_result = pd.concat(results).reset_index(drop=True)
            final_result['year'] = pd.to_datetime(final_result['registrationdate']).dt.year
            final_result = final_result.sort_values(by='year', ascending=True).drop(columns='year')

            st.write("Property Details:")
            st.dataframe(final_result)

            ownership_history = []
            missing_links = []
            previous_owner = None

            for i, row in final_result.iterrows():
                ownership_history.append({
                    "owner_name": row['purchaserparty'],
                    "acquisition_date": row['registrationdate'],
                    "survey_number": row['docno'],
                    "area": row['areaname'],
                    "property_location": row['propertydescription'],
                    "sale_price": row['consideration_amt']
                })
                if previous_owner and previous_owner != row['sellerparty']:
                    missing_links.append(f"Ownership transferred from {previous_owner} to {row['sellerparty']} without a recorded transaction.")
                previous_owner = row['purchaserparty']

            prompt = create_prompt_maharashtra(str_date, ownership_history, missing_links)
            api_key = "AIzaSyApU6L3bNMD0bxyLtay2yn7S4yMatskIpI"  # Replace with your actual API key
            report_content = generate_report(prompt, api_key)

            if report_content:
                sanitized_property_desc = sanitize_filename(property_description)
                save_to_word(report_content, f"{sanitized_property_desc}_search_title_report", "ADVOCATE PRAVIN D SANGVIKAR", "SAMADHAN, N-7, R-28, H-15/28, MHADA, H.I.G.,CIDCO, AURANGABAD.")
            else:
                st.error("Failed to generate the Search Title Report. Please check the API response.")
        else:
            st.error("No matching property details found in the uploaded files.")
